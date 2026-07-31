from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "owner-guide" / "CRM-Workflow-and-Functions-Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5F6B76"
HEADER_FILL = "E8EEF5"
BORDER = "B9C4CF"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=10, color=INK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=9.5, color="000000")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=11, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11)
    return paragraph


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("ALLEATO CRM  |  WORKFLOW AND FUNCTION GUIDE")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7200, 2160])
    table.rows[0].cells[0].text = "Alleato Group - Internal operating reference"
    table.rows[0].cells[1].text = "July 2026"
    for index, cell in enumerate(table.rows[0].cells):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            set_run_font(run, size=8, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(section)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("Alleato CRM Workflow and Function Guide")
    set_run_font(run, size=25, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "How relationship management, opportunities, existing tasks, communications, documents, and project conversion work together"
    )
    set_run_font(run, size=12.5, color=MUTED)
    add_table(
        doc,
        ["Audience", "System", "Updated"],
        [["Business development, project leadership, CRM administrators", "Alleato Project Management CRM", "July 29, 2026"]],
        [4320, 2880, 2160],
    )

    doc.add_heading("1. Purpose", level=1)
    add_body(
        doc,
        "The CRM is the relationship and opportunity layer inside the existing Alleato Project Management application. It uses the Company Directory as the source of truth for companies, the existing Tasks system for follow-ups, and existing document and project records for supporting context. It does not create a second company list or a separate task application.",
    )

    doc.add_heading("2. Core workflow", level=1)
    for item in (
        "Open a company in the Company Directory.",
        "Select Add to CRM to enroll the existing company as a relationship.",
        "Review the owner, lifecycle stage, health, activity, open deals, and next follow-up.",
        "Record calls, emails, meetings, and notes in relationship history.",
        "Create a deal and move it through the sales pipeline.",
        "Create follow-ups from the deal. They appear in the existing Tasks system and on the CRM deal.",
        "Review communication matching suggestions before accepting them into CRM history.",
        "For a won deal, create the project and let the scheduled reconciliation confirm Acumatica status.",
    ):
        add_number(doc, item)

    doc.add_heading("3. CRM workspace", level=1)
    add_table(
        doc,
        ["Area", "Purpose", "Expected behavior"],
        [
            ("Relationships", "Daily work queue", "Search and filter enrolled companies; review health, open a company, archive, or restore."),
            ("Pipeline", "Stage view", "Scan deals by stage and move an opportunity to another valid stage."),
            ("Deals", "Opportunity list", "Create, filter, open, edit, archive, restore, and convert deals."),
            ("Activity", "Relationship history", "Log, edit, and delete manual activity; view accepted communication."),
            ("Matching", "Review queue", "Accept or reject suggested communication-to-company matches."),
            ("Settings", "Operating rules", "Manage health thresholds, stale-deal timing, timezone, and matching behavior."),
        ],
        [1620, 2520, 5220],
    )

    doc.add_heading("4. Company enrollment and ownership", level=1)
    add_body(
        doc,
        "The Company Directory remains authoritative for the company name and ERP-owned data. Adding a company to CRM creates only a relationship profile containing the CRM owner, lifecycle, health, reason, last meaningful activity, and archive state.",
    )
    add_body(
        doc,
        "Owners and administrators can update relationship records. CRM access is controlled by the CRM permission module; application administrators receive administrative CRM access.",
    )

    doc.add_heading("5. Relationship health", level=1)
    add_body(doc, "Health is derived from meaningful activity and configured thresholds:")
    for item in (
        "Healthy - recent meaningful activity is inside the active threshold.",
        "Watch - activity is older than the active threshold but inside the watch threshold.",
        "Attention - activity is older than the watch threshold or a follow-up/deal condition needs attention.",
    ):
        add_bullet(doc, item)
    add_body(doc, "Health supports decisions; it does not replace judgment. The displayed reason explains the status.")

    doc.add_heading("6. Deals and pipeline", level=1)
    add_body(
        doc,
        "A deal belongs to one enrolled company, one owner, one pipeline, and one stage. The standard flow is Lead -> Qualified -> Proposal / Bid -> Negotiation -> Won or Lost.",
    )
    add_body(
        doc,
        "Stage changes use the saved row version to prevent one user from overwriting a newer change. Lost opportunities require a reason. A won deal with a linked project cannot be reopened until the project link is deliberately removed with a change reason.",
    )

    doc.add_heading("7. Follow-ups use the existing Tasks system", level=1)
    add_body(doc, "CRM follow-ups are ordinary records in the existing Tasks system, not a separate CRM-only list.")
    for item in (
        "The task is linked to the company and deal.",
        "It is assigned to the relationship owner.",
        "It appears in the existing Tasks experience.",
        "Its source is labeled CRM and links back to the CRM record.",
        "Completing it from Tasks or CRM updates the same record.",
    ):
        add_bullet(doc, item)
    add_body(
        doc,
        "Recommended daily practice: begin in Tasks, complete due calls and outreach, then return to CRM to record the outcome and advance the opportunity.",
    )

    doc.add_heading("8. Communication matching and email behavior", level=1)
    add_body(
        doc,
        "The CRM does not directly send, delete, or modify email. It reviews communication metadata already ingested by the application's Outlook, Teams, and Fireflies source pipelines.",
    )
    for item in (
        "Read recent eligible communication metadata.",
        "Exclude private, restricted, and leadership-only sources.",
        "Match an exact business email domain or company-name signal with enrolled CRM companies.",
        "Create a pending suggestion with a confidence score.",
        "Wait for a user to accept or reject the suggestion.",
    ):
        add_number(doc, item)
    add_body(doc, "Accepted suggestions become CRM activity. Source systems remain read-only.")

    doc.add_heading("9. AI status", level=1)
    add_body(
        doc,
        "This release includes automated, confidence-based communication matching with human review. It is deterministic and does not autonomously send messages, change deals, or complete tasks.",
    )
    add_body(
        doc,
        "A future generative CRM assistant should use the application's existing AI framework and remain permission-aware, source-cited, and confirmation-gated for every write action.",
    )

    doc.add_heading("10. Documents", level=1)
    add_body(
        doc,
        "A deal can link to an existing document record. Linking does not copy the file or change source permissions. Removing the CRM link does not delete the underlying document.",
    )

    doc.add_heading("11. Won-deal project conversion", level=1)
    for item in (
        "Select Create project on a won deal without a project.",
        "The CRM calls the existing project creation workflow with an idempotency key.",
        "The new project is linked to the deal and the conversion waits for ERP reconciliation.",
        "The scheduled CRM job checks the linked project.",
        "When the project has an Acumatica project identifier, the conversion is marked synchronized.",
    ):
        add_number(doc, item)
    add_body(
        doc,
        "The CRM never fabricates an Acumatica success state. Failed attempts remain available for diagnosis and retry.",
    )

    doc.add_heading("12. Archive and recovery", level=1)
    add_body(
        doc,
        "Archiving is reversible and requires a reason. Open dependencies can prevent an account or deal from being archived. Project links require an explicit removal reason. The application shows server-confirmed success only after the write completes.",
    )

    doc.add_heading("13. Suggested daily operating rhythm", level=1)
    for item in (
        "Review Tasks for due and overdue CRM follow-ups.",
        "Open CRM > Relationships and filter for Attention or Watch.",
        "Complete outreach and log the result.",
        "Review CRM > Matching and accept or reject communication suggestions.",
        "Open CRM > Pipeline and advance opportunities whose next step is complete.",
        "Review won deals ready for project creation.",
        "Use CRM > Settings only for administrative rule changes.",
    ):
        add_number(doc, item)

    doc.add_heading("14. Data ownership summary", level=1)
    add_table(
        doc,
        ["Record", "System of record"],
        [
            ("Company identity and ERP data", "Company Directory / ERP integration"),
            ("CRM lifecycle, owner, health, deals, activity", "CRM tables"),
            ("Follow-up work", "Existing Tasks system"),
            ("Files and source permissions", "Existing document system"),
            ("Created project", "Existing Projects system"),
            ("Acumatica synchronization state", "Project integration and CRM conversion record"),
            ("Outlook, Teams, Fireflies content", "Existing source-ingestion pipelines; CRM reads eligible metadata only"),
        ],
        [3420, 5940],
    )

    doc.core_properties.title = "Alleato CRM Workflow and Function Guide"
    doc.core_properties.subject = "CRM operating workflow and function reference"
    doc.core_properties.author = "Alleato Group"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
