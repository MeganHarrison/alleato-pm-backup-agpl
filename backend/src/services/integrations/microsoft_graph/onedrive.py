"""
OneDrive / SharePoint File Ingestion
Fetches and extracts text from project documents.
"""
import hashlib
import io
import logging
import os
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from datetime import date, datetime, time, timezone
from pathlib import PurePosixPath
from typing import Optional
from urllib.parse import quote

from src.services.ingestion.project_assignment import (
    AssignmentTarget,
    ProjectAssigner,
)

from ...supabase_helpers import (
    SupabaseRagStore,
    get_rag_read_client,
    storage_upload_with_retry,
)
from .client import get_graph_client
from .project_documents import (
    DOCUMENT_BUCKET,
    graph_id_safe,
    metadata_text_storage,
    project_document_payload_from_graph_item,
    source_path,
    upsert_project_document_by_source,
)
from .project_inference import infer_assignment_target

logger = logging.getLogger(__name__)
_WARNED_OPTIONAL_DEPENDENCIES: set[str] = set()
_STREAMED_PDF_MAX_BYTES = 500 * 1024 * 1024


@dataclass(frozen=True)
class SharePointSyncResult:
    """The outcome of one SharePoint delta pass.

    `retry_required` deliberately keeps the existing delta cursor in place.
    Otherwise a file that fails after Graph has supplied its delta entry is
    silently skipped forever when the new cursor is saved.
    """

    items_synced: int
    delta_token: str
    items_failed: int = 0
    items_unchanged: int = 0
    items_excluded: int = 0
    excluded_by_extension: dict[str, int] | None = None
    retry_required: bool = False
    retry_reason: Optional[str] = None
    inventory_complete: bool = True


def _warn_once(key: str, message: str, *args) -> None:
    if key in _WARNED_OPTIONAL_DEPENDENCIES:
        logger.debug(message, *args)
        return
    _WARNED_OPTIONAL_DEPENDENCIES.add(key)
    logger.warning(message, *args)


def _bounded_int_env(name: str, default: int, minimum: int, maximum: int) -> int:
    raw = os.environ.get(name)
    try:
        value = int(raw) if raw not in (None, "") else default
    except (TypeError, ValueError):
        return default
    return max(minimum, min(maximum, value))


def _actual_parent_path(item: dict) -> str:
    """Extract the real folder path from Graph API parentReference.
    Graph returns paths like '/drive/root:/Alleato Group/2026 Jobs/ProjectX'."""
    ref_path = (item.get("parentReference") or {}).get("path", "")
    if "root:/" in ref_path:
        return ref_path.split("root:/", 1)[1]
    return ref_path.lstrip("/")


def _item_source_path(item: dict, folder_path: str, name: str) -> str:
    """Build source_path using the actual parent folder from the Graph item,
    preserving subfolders that the configured root folder misses."""
    actual_parent = _actual_parent_path(item)
    if actual_parent:
        return str(PurePosixPath(actual_parent) / name)
    return source_path(folder_path, name)


def _project_subfolder(item: dict, root_folder: str) -> Optional[str]:
    """Return the first subfolder name below root_folder, if any.
    e.g. root='Alleato Group/2026 Jobs', actual parent='Alleato Group/2026 Jobs/Vermillion Rise'
    → 'Vermillion Rise'"""
    actual_parent = _actual_parent_path(item)
    root = root_folder.strip("/")
    parent = actual_parent.strip("/")
    if root and parent.startswith(root + "/"):
        parts = [p for p in parent[len(root) + 1:].split("/") if p]
        return parts[0] if parts else None
    return None


def _strip_folder_prefix(folder_name: str) -> str:
    """Strip numeric job-number prefix like '25- 104 ' or '26-001 ' from folder names."""
    import re
    stripped = re.sub(r"^\d{2}-\s*\d+\s+", "", folder_name).strip()
    return stripped if stripped else folder_name


def _normalize_job_number(value: str | None) -> str | None:
    match = re.search(r"\b(\d{2})\s*-\s*(\d{3})\b", str(value or ""))
    return f"{match.group(1)}-{match.group(2)}" if match else None


def _job_number_from_source_path(item: dict, root_folder: str) -> str | None:
    values = [_actual_parent_path(item), root_folder]
    for value in values:
        for part in str(value or "").split("/"):
            job_number = _normalize_job_number(part)
            if job_number:
                return job_number
    return None


def _lookup_project_by_job_number(
    supabase_client,
    job_number: str,
) -> Optional[int]:
    """Prefer an exact canonical project-number match over fuzzy title inference."""
    normalized = _normalize_job_number(job_number)
    if not normalized:
        return None
    candidates: dict[int, dict] = {}
    for column in ("project_number", "job number"):
        try:
            response = (
                supabase_client.from_("projects")
                .select('id,is_development,archived,project_number,"job number"')
                .eq(column, normalized)
                .execute()
            )
            for row in response.data or []:
                if row.get("id") is not None:
                    candidates[int(row["id"])] = row
        except Exception as exc:
            logger.warning(
                "[OneDrive] project lookup by %s=%s failed: %s",
                column,
                normalized,
                exc,
            )
    if not candidates:
        return None
    ordered = sorted(
        candidates.values(),
        key=lambda row: (
            bool(row.get("archived")),
            bool(row.get("is_development")),
            int(row.get("id") or 0),
        ),
    )
    return int(ordered[0]["id"])


def _lookup_project_by_folder(supabase_client, folder_name: str) -> Optional[int]:
    """Case-insensitive project name match with fallback for numeric-prefix folder names."""
    try:
        res = (
            supabase_client.from_("projects")
            .select("id")
            .ilike("name", folder_name)
            .limit(1)
            .execute()
        )
        if res.data:
            return int(res.data[0]["id"])
    except Exception as exc:
        logger.warning("[OneDrive] folder→project lookup failed: %s", exc)
        return None

    # Strip numeric job-number prefix (e.g. "25- 104 Danville Theatre" → "Danville Theatre")
    # then do a partial contains match
    stripped = _strip_folder_prefix(folder_name)
    if stripped and stripped != folder_name:
        try:
            res = (
                supabase_client.from_("projects")
                .select("id")
                .ilike("name", f"%{stripped}%")
                .limit(1)
                .execute()
            )
            if res.data:
                logger.info("[OneDrive] folder '%s' matched project via stripped name '%s'", folder_name, stripped)
                return int(res.data[0]["id"])
        except Exception as exc:
            logger.warning("[OneDrive] stripped folder→project lookup failed: %s", exc)
    return None


def _target_for_matched_project(
    supabase_client,
    *,
    project_id: int,
    method: str,
) -> AssignmentTarget:
    target = ProjectAssigner(supabase_client).assign_scope(
        meeting_title="",
        participants=[],
        existing_project_id=project_id,
        migrate_mapped_existing=True,
    )
    return AssignmentTarget(
        project_id=target.project_id,
        business_area_id=target.business_area_id,
        legacy_project_id=target.legacy_project_id,
        method=method,
        confidence=1.0,
    )


def _assign_scope(
    supabase_client,
    item: dict,
    root_folder: str,
    title: str,
    content: str,
    participants: list,
    expected_project_number: str | None = None,
) -> AssignmentTarget:
    """Try exact source matches first, then infer one typed destination."""
    job_number = (
        _normalize_job_number(expected_project_number)
        or _job_number_from_source_path(item, root_folder)
    )
    if job_number:
        project_id = _lookup_project_by_job_number(supabase_client, job_number)
        if project_id:
            logger.info(
                "[OneDrive] Assigned project_id=%s via exact job number %s",
                project_id,
                job_number,
            )
            return _target_for_matched_project(
                supabase_client,
                project_id=project_id,
                method="job_number",
            )
    subfolder = _project_subfolder(item, root_folder)
    if subfolder:
        project_id = _lookup_project_by_folder(supabase_client, subfolder)
        if project_id:
            logger.info("[OneDrive] Assigned project_id=%s via folder name '%s'", project_id, subfolder)
            return _target_for_matched_project(
                supabase_client,
                project_id=project_id,
                method="folder_name",
            )
    return _infer_assignment_scope(
        supabase_client,
        title=title,
        content=content,
        participants=participants,
    )


def _assignment_tag(target: AssignmentTarget) -> str:
    if target.business_area_id is not None:
        return f"business_area_auto:{target.method}"
    if target.project_id is not None:
        return f"project_auto:{target.method}"
    return "unassigned"


def _infer_assignment_scope(
    supabase_client,
    *,
    title: str,
    content: str,
    participants: list,
) -> AssignmentTarget:
    return infer_assignment_target(
        supabase_client,
        title=title,
        content=content,
        participants=participants,
    )


def _set_document_scope(
    supabase_client,
    document_id: str,
    target: AssignmentTarget,
) -> None:
    store = SupabaseRagStore(supabase_client)
    if hasattr(store, "set_document_scope"):
        store.set_document_scope(
            document_id,
            project_id=target.project_id,
            business_area_id=target.business_area_id,
        )
        return
    # Lightweight offline stores used by focused ingestion tests do not mirror
    # the separate RAG database, but still exercise the exact app payload.
    supabase_client.from_("document_metadata").update(
        {
            "project_id": target.project_id,
            "business_area_id": target.business_area_id,
        }
    ).eq("id", document_id).execute()

# File types the canonical Graph ingestion path can materialize as text.
# A newly added extension also requires a deliberate historical cursor replay:
# Graph delta cursors do not resend files skipped before support was added.
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".md",
    ".csv",
    ".eml",
    ".html",
    ".mpp",
    ".msg",
    ".pptx",
    ".rtf",
    ".xls",
    ".xlsx",
    ".xlsm",
}
# These formats do not contain directly vectorizable text. They remain visible
# in the per-scope receipt with an exact extension/count and are not represented
# as searchable text. A new extension is never silently added here.
GOVERNED_NON_TEXT_EXTENSIONS = {
    ".avif",
    ".bak",
    ".ctb",
    ".dng",
    ".download",
    ".dwg",
    ".dwl",
    ".dwl2",
    ".geprint",
    ".heic",
    ".hlf",
    ".hvuf",
    ".insv",
    ".jpeg",
    ".jpg",
    ".kss",
    ".layout",
    ".lbl",
    ".lnk",
    ".log",
    ".lrv",
    ".mov",
    ".mp4",
    ".pc3",
    ".png",
    ".rvt",
    ".skb",
    ".skp",
    ".svg",
    ".url",
    ".wxf",
    ".xml",
    ".zip",
}
# These formats can contain decision-grade text but do not yet have a complete
# extractor in the canonical path. They therefore fail the scope and preserve
# the prior cursor instead of being mislabeled as a valid exclusion.
TEXT_BEARING_UNSUPPORTED_EXTENSIONS = {
}
DEFAULT_MAX_FILE_SIZE_MB = 250
HARD_MAX_FILE_SIZE_MB = 500


def _max_file_size_bytes() -> int:
    raw = os.getenv(
        "GRAPH_INGEST_MAX_FILE_SIZE_MB",
        str(DEFAULT_MAX_FILE_SIZE_MB),
    )
    try:
        configured_mb = int(raw)
    except ValueError as exc:
        raise RuntimeError(
            "GRAPH_INGEST_MAX_FILE_SIZE_MB must be an integer"
        ) from exc
    if configured_mb < 1 or configured_mb > HARD_MAX_FILE_SIZE_MB:
        raise RuntimeError(
            "GRAPH_INGEST_MAX_FILE_SIZE_MB must be between 1 and "
            f"{HARD_MAX_FILE_SIZE_MB}"
        )
    return configured_mb * 1024 * 1024


def _fetch_rag_document_text(doc_id: str) -> str:
    try:
        row = (
            get_rag_read_client()
            .from_("rag_document_metadata")
            .select("content,raw_text")
            .eq("id", doc_id)
            .single()
            .execute()
            .data
            or {}
        )
        return str(row.get("content") or row.get("raw_text") or "")
    except Exception:
        return ""


def _extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using pypdf with a PyMuPDF fallback."""
    pypdf_error: Exception | None = None
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        extracted = "\n\n".join(pages)
        if len(extracted.strip()) >= 50:
            return extracted
    except ImportError:
        _warn_once("pypdf_missing", "[OneDrive] pypdf not installed — skipping PDF extraction")
    except Exception as e:
        pypdf_error = e
        logger.warning("[OneDrive] pypdf PDF extraction failed, trying PyMuPDF fallback: %s", e)

    try:
        import fitz

        with fitz.open(stream=content, filetype="pdf") as document:
            pages = []
            for page in document:
                text = page.get_text("text") or ""
                if text.strip():
                    pages.append(text)
            extracted = "\n\n".join(pages)
            if extracted.strip():
                return extracted
    except ImportError:
        _warn_once("pymupdf_missing", "[OneDrive] PyMuPDF not installed — PDF fallback unavailable")
    except Exception as e:
        logger.warning("[OneDrive] PyMuPDF PDF extraction failed: %s", e)

    if pypdf_error:
        logger.warning("[OneDrive] PDF extraction failed after fallback: %s", pypdf_error)
    return ""


def _extract_text_from_docx(content: bytes) -> str:
    """Extract paragraphs, tables, headers, and footers from DOCX bytes."""
    try:
        import docx
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = docx.Document(io.BytesIO(content))
        parts: list[str] = []
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                text = Paragraph(child, doc).text.strip()
                if text:
                    parts.append(text)
            elif isinstance(child, CT_Tbl):
                table = Table(child, doc)
                parts.append("[Table]")
                for row_index, row in enumerate(table.rows, start=1):
                    for column_index, cell in enumerate(row.cells, start=1):
                        cell_text = "\n".join(
                            paragraph.text.strip()
                            for paragraph in cell.paragraphs
                            if paragraph.text.strip()
                        )
                        if cell_text:
                            parts.append(f"R{row_index}C{column_index}: {cell_text}")

        seen_section_text: set[tuple[str, str]] = set()
        for section in doc.sections:
            for label, container in (
                ("Header", section.header),
                ("First-page header", section.first_page_header),
                ("Even-page header", section.even_page_header),
                ("Footer", section.footer),
                ("First-page footer", section.first_page_footer),
                ("Even-page footer", section.even_page_footer),
            ):
                text = "\n".join(
                    paragraph.text.strip()
                    for paragraph in container.paragraphs
                    if paragraph.text.strip()
                )
                identity = (label, text)
                if text and identity not in seen_section_text:
                    parts.append(f"[{label}]\n{text}")
                    seen_section_text.add(identity)
        return "\n".join(parts)
    except ImportError:
        logger.warning("[OneDrive] python-docx not installed — skipping DOCX extraction")
        return _extract_text_from_docx_archive(content)
    except Exception as e:
        logger.warning(
            "[OneDrive] python-docx extraction failed, trying direct OpenXML fallback: %s",
            e,
        )
        return _extract_text_from_docx_archive(content)


def _extract_text_from_docx_archive(content: bytes) -> str:
    """Read text-bearing Word OpenXML parts without trusting package relations.

    Some SharePoint files contain a valid ``word/document.xml`` but a malformed
    root relationship that points python-docx at a theme-manager part. Reading
    the governed text parts directly preserves the document body, headers,
    footers, comments, and notes instead of silently losing the entire file.
    """
    import zipfile
    from xml.etree import ElementTree

    text_part = re.compile(
        r"^word/(?:document|comments|footnotes|endnotes|"
        r"header\d+|footer\d+)\.xml$"
    )
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            names = sorted(name for name in archive.namelist() if text_part.match(name))
            if "word/document.xml" not in names:
                raise RuntimeError("word/document.xml is missing")
            lines: list[str] = []
            for name in names:
                root = ElementTree.fromstring(archive.read(name))
                texts = [
                    str(node.text).strip()
                    for node in root.iter()
                    if node.tag.endswith("}t")
                    and node.text
                    and str(node.text).strip()
                ]
                if texts:
                    lines.append(f"[{name}]")
                    lines.extend(texts)
            return "\n".join(lines).strip()
    except Exception as exc:
        raise RuntimeError(f"DOCX OpenXML fallback failed: {exc}") from exc


def _docx_embedded_image_names(content: bytes) -> list[str]:
    """Return the governed image parts that require OCR in an image-only DOCX."""
    import zipfile

    supported = {".bmp", ".heif", ".jpeg", ".jpg", ".png", ".tif", ".tiff"}
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            return sorted(
                name
                for name in archive.namelist()
                if name.startswith("word/media/")
                and PurePosixPath(name).suffix.lower() in supported
            )
    except Exception as exc:
        raise RuntimeError(f"DOCX embedded-image inventory failed: {exc}") from exc


def _workbook_cell_text(value: object) -> str:
    """Return a deterministic, readable representation of a workbook cell."""
    if isinstance(value, (datetime, date, time)):
        return value.isoformat()
    if isinstance(value, bytes):
        return value.decode("utf-8", errors="replace")
    return str(value).replace("\x00", "")


def _extract_text_from_openxml_parts(content: bytes) -> str:
    """Read workbook cells directly when third-party style parsing is invalid.

    Excel can open workbooks whose named-style references are inconsistent,
    while openpyxl refuses the entire file before exposing any cells. The
    fallback intentionally ignores presentation styles and follows the Open XML
    workbook relationships, shared strings, formulas, inline strings, and raw
    cached values so source content remains complete and attributable.
    """
    import posixpath
    import zipfile
    from xml.etree import ElementTree

    spreadsheet_ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    relationships_ns = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    )
    package_rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ns = {"s": spreadsheet_ns, "r": relationships_ns}

    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        shared_strings: list[str] = []
        if "xl/sharedStrings.xml" in archive.namelist():
            shared_root = ElementTree.fromstring(
                archive.read("xl/sharedStrings.xml")
            )
            for item in shared_root.findall("s:si", ns):
                shared_strings.append(
                    "".join(
                        str(node.text or "")
                        for node in item.iter(f"{{{spreadsheet_ns}}}t")
                    )
                )

        workbook_root = ElementTree.fromstring(archive.read("xl/workbook.xml"))
        relation_root = ElementTree.fromstring(
            archive.read("xl/_rels/workbook.xml.rels")
        )
        relation_targets = {
            str(node.attrib.get("Id") or ""): str(node.attrib.get("Target") or "")
            for node in relation_root.findall(f"{{{package_rel_ns}}}Relationship")
        }

        lines: list[str] = []
        for sheet in workbook_root.findall("s:sheets/s:sheet", ns):
            sheet_name = str(sheet.attrib.get("name") or "Worksheet")
            relation_id = str(
                sheet.attrib.get(f"{{{relationships_ns}}}id") or ""
            )
            target = relation_targets.get(relation_id)
            if not target:
                raise RuntimeError(
                    f"XLSX worksheet {sheet_name!r} has no relationship target"
                )
            normalized_target = posixpath.normpath(
                posixpath.join("xl", target)
                if not target.startswith("/")
                else target.lstrip("/")
            )
            sheet_root = ElementTree.fromstring(archive.read(normalized_target))
            lines.append(f"[Worksheet: {sheet_name}]")

            for cell in sheet_root.findall(".//s:sheetData/s:row/s:c", ns):
                coordinate = str(cell.attrib.get("r") or "cell")
                cell_type = str(cell.attrib.get("t") or "")
                formula = cell.find("s:f", ns)
                raw_value = cell.find("s:v", ns)
                value: object | None = (
                    raw_value.text if raw_value is not None else None
                )

                if cell_type == "s" and value not in (None, ""):
                    try:
                        value = shared_strings[int(str(value))]
                    except (ValueError, IndexError) as exc:
                        raise RuntimeError(
                            f"XLSX shared-string index is invalid at {coordinate}"
                        ) from exc
                elif cell_type == "inlineStr":
                    value = "".join(
                        str(node.text or "")
                        for node in cell.iter(f"{{{spreadsheet_ns}}}t")
                    )
                elif cell_type == "b" and value is not None:
                    value = "TRUE" if str(value) == "1" else "FALSE"

                if formula is not None:
                    formula_text = str(formula.text or "")
                    if not formula_text.startswith("="):
                        formula_text = f"={formula_text}"
                    if value in (None, ""):
                        lines.append(
                            f"{coordinate} formula: {formula_text}"
                        )
                    else:
                        lines.append(
                            f"{coordinate} formula: {formula_text} | cached value: "
                            f"{_workbook_cell_text(value)}"
                        )
                elif value not in (None, ""):
                    lines.append(
                        f"{coordinate}: {_workbook_cell_text(value)}"
                    )
            lines.append("")

        return "\n".join(lines).strip()


def _extract_text_from_openxml_workbook(content: bytes) -> str:
    """Materialize every non-empty XLSX/XLSM cell with worksheet provenance.

    The formula workbook preserves expressions while the data-only workbook
    supplies Excel's last cached result when one exists. No row, worksheet, or
    character cap is applied here; the downstream chunker owns model-sized
    segmentation of the complete materialized text.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for XLSX/XLSM extraction") from exc

    formula_workbook = None
    value_workbook = None
    try:
        formula_workbook = load_workbook(
            io.BytesIO(content), read_only=True, data_only=False, keep_links=True
        )
        value_workbook = load_workbook(
            io.BytesIO(content), read_only=True, data_only=True, keep_links=True
        )

        lines: list[str] = []
        for formula_sheet in formula_workbook.worksheets:
            value_sheet = value_workbook[formula_sheet.title]
            lines.append(f"[Worksheet: {formula_sheet.title}]")
            for formula_row, value_row in zip(
                formula_sheet.iter_rows(), value_sheet.iter_rows()
            ):
                for formula_cell, value_cell in zip(formula_row, value_row):
                    formula_value = formula_cell.value
                    cached_value = value_cell.value
                    if formula_value is None and cached_value is None:
                        continue

                    coordinate = formula_cell.coordinate
                    if formula_cell.data_type == "f":
                        formula_text = _workbook_cell_text(formula_value)
                        if cached_value is None:
                            lines.append(f"{coordinate} formula: {formula_text}")
                        else:
                            lines.append(
                                f"{coordinate} formula: {formula_text} | cached value: "
                                f"{_workbook_cell_text(cached_value)}"
                            )
                    else:
                        lines.append(f"{coordinate}: {_workbook_cell_text(formula_value)}")
            lines.append("")
        return "\n".join(lines).strip()
    except Exception as exc:
        logger.warning(
            "[OneDrive] openpyxl workbook parsing failed; retrying direct Open "
            "XML cell extraction: %s",
            exc,
        )
        try:
            return _extract_text_from_openxml_parts(content)
        except Exception as fallback_exc:
            raise RuntimeError(
                "XLSX/XLSM extraction failed in both openpyxl and direct "
                f"Open XML paths: {fallback_exc}"
            ) from fallback_exc
    finally:
        if formula_workbook is not None:
            formula_workbook.close()
        if value_workbook is not None:
            value_workbook.close()


def _extract_text_from_legacy_doc(content: bytes) -> str:
    """Extract complete text from an OLE Word ``.doc`` using antiword."""
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc") as source:
            source.write(content)
            source.flush()
            result = subprocess.run(
                ["antiword", source.name],
                capture_output=True,
                check=False,
                timeout=60,
            )
    except FileNotFoundError as exc:
        raise RuntimeError(
            "antiword is required for legacy Word .doc extraction"
        ) from exc
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError("legacy Word .doc extraction timed out") from exc

    if result.returncode != 0:
        detail = result.stderr.decode("utf-8", errors="replace").strip()
        raise RuntimeError(
            "legacy Word .doc extraction failed"
            + (f": {detail}" if detail else "")
        )
    return result.stdout.decode("utf-8", errors="replace").strip()


def _extract_text_from_xls(content: bytes) -> str:
    """Materialize every non-empty legacy XLS cell with sheet provenance."""
    try:
        import xlrd
    except ImportError as exc:
        raise RuntimeError("xlrd is required for XLS extraction") from exc
    try:
        workbook = xlrd.open_workbook(file_contents=content, formatting_info=False)
        lines: list[str] = []
        for sheet in workbook.sheets():
            lines.append(f"[Worksheet: {sheet.name}]")
            for row_index in range(sheet.nrows):
                for column_index in range(sheet.ncols):
                    cell = sheet.cell(row_index, column_index)
                    if cell.value in (None, ""):
                        continue
                    coordinate = f"R{row_index + 1}C{column_index + 1}"
                    lines.append(f"{coordinate}: {_workbook_cell_text(cell.value)}")
            lines.append("")
        return "\n".join(lines).strip()
    except Exception as exc:
        raise RuntimeError(f"XLS extraction failed: {exc}") from exc


def _extract_text_from_pptx(content: bytes) -> str:
    """Extract all DrawingML text from slides, notes, comments, and masters."""
    import zipfile
    from xml.etree import ElementTree

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("ppt/") and name.endswith(".xml")
            )
            lines: list[str] = []
            for name in names:
                root = ElementTree.fromstring(archive.read(name))
                texts = [
                    str(node.text).strip()
                    for node in root.iter()
                    if node.tag.endswith("}t")
                    and node.text
                    and str(node.text).strip()
                ]
                if texts:
                    lines.append(f"[{name}]")
                    lines.extend(texts)
            return "\n".join(lines).strip()
    except Exception as exc:
        raise RuntimeError(f"PPTX extraction failed: {exc}") from exc


def _extract_text_from_html(content: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("beautifulsoup4 is required for HTML extraction") from exc
    try:
        soup = BeautifulSoup(content, "html.parser")
        for node in soup(["script", "style", "noscript"]):
            node.decompose()
        return "\n".join(
            line.strip()
            for line in soup.get_text("\n").splitlines()
            if line.strip()
        )
    except Exception as exc:
        raise RuntimeError(f"HTML extraction failed: {exc}") from exc


def _extract_text_from_eml(content: bytes) -> str:
    from email import policy
    from email.parser import BytesParser

    try:
        message = BytesParser(policy=policy.default).parsebytes(content)
        lines = [
            f"Subject: {message.get('subject', '')}",
            f"From: {message.get('from', '')}",
            f"To: {message.get('to', '')}",
            f"Cc: {message.get('cc', '')}",
            f"Date: {message.get('date', '')}",
        ]
        parts = list(message.walk()) if message.is_multipart() else [message]
        for part in parts:
            if part.get_content_disposition() == "attachment":
                continue
            content_type = part.get_content_type()
            if content_type not in {"text/plain", "text/html"}:
                continue
            payload = part.get_payload(decode=True)
            if payload is None:
                body = str(part.get_payload() or "")
            else:
                charset = part.get_content_charset() or "utf-8"
                body = payload.decode(charset, errors="replace")
            if content_type == "text/html":
                body = _extract_text_from_html(body.encode("utf-8"))
            if body.strip():
                lines.append(body.strip())
        return "\n".join(lines).strip()
    except Exception as exc:
        raise RuntimeError(f"EML extraction failed: {exc}") from exc


def _extract_text_from_rtf(content: bytes) -> str:
    """Decode RTF control words, escaped bytes, and Unicode text."""
    text = content.decode("latin-1", errors="replace")
    text = re.sub(
        r"\\'([0-9a-fA-F]{2})",
        lambda match: bytes.fromhex(match.group(1)).decode(
            "cp1252",
            errors="replace",
        ),
        text,
    )
    text = re.sub(
        r"\\u(-?\d+)\??",
        lambda match: chr(int(match.group(1)) % 65536),
        text,
    )
    text = re.sub(r"\\(par|line)\b ?", "\n", text)
    text = re.sub(r"\\tab\b ?", "\t", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?|\\[{}\\]", "", text)
    text = text.replace("{", "").replace("}", "")
    return "\n".join(
        line.strip() for line in text.splitlines() if line.strip()
    )


def _extract_text_from_msg(content: bytes) -> str:
    """Extract Outlook MSG headers and body with the maintained parser."""
    import tempfile

    try:
        import extract_msg
    except ImportError as exc:
        raise RuntimeError("extract-msg is required for MSG extraction") from exc
    try:
        with tempfile.NamedTemporaryFile(suffix=".msg") as handle:
            handle.write(content)
            handle.flush()
            message = extract_msg.openMsg(handle.name)
            try:
                lines = [
                    f"Subject: {message.subject or ''}",
                    f"From: {message.sender or ''}",
                    f"To: {message.to or ''}",
                    f"Cc: {message.cc or ''}",
                    f"Date: {message.date or ''}",
                    str(message.body or ""),
                ]
                return "\n".join(line for line in lines if line.strip()).strip()
            finally:
                message.close()
    except Exception as exc:
        raise RuntimeError(f"MSG extraction failed: {exc}") from exc


def _extract_text_from_mpp(content: bytes) -> str:
    """Materialize an MPP in an isolated, time-bounded MPXJ process."""
    try:
        timeout_seconds = int(
            os.environ.get("MPP_EXTRACTION_TIMEOUT_SECONDS", "90")
        )
    except ValueError:
        timeout_seconds = 90
    timeout_seconds = max(10, min(timeout_seconds, 300))
    try:
        max_output_bytes = int(
            os.environ.get(
                "MPP_EXTRACTION_MAX_OUTPUT_BYTES",
                str(20 * 1024 * 1024),
            )
        )
    except ValueError:
        max_output_bytes = 20 * 1024 * 1024
    max_output_bytes = max(1024, min(max_output_bytes, 100 * 1024 * 1024))

    worker = os.path.join(os.path.dirname(__file__), "mpp_worker.py")
    with tempfile.NamedTemporaryFile(suffix=".mpp") as handle:
        handle.write(content)
        handle.flush()
        try:
            result = subprocess.run(
                [sys.executable, worker, handle.name],
                capture_output=True,
                text=True,
                check=False,
                timeout=timeout_seconds,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError(
                "MPP extraction exceeded the governed "
                f"{timeout_seconds}-second timeout"
            ) from exc
    if result.returncode != 0:
        reason = (result.stderr or result.stdout or "unknown worker error").strip()
        raise RuntimeError(f"MPP extraction failed: {reason[:1000]}")
    encoded_size = len(result.stdout.encode("utf-8"))
    if encoded_size > max_output_bytes:
        raise RuntimeError(
            "MPP extraction output exceeded the governed "
            f"{max_output_bytes}-byte limit"
        )
    return result.stdout.strip()


def _extract_text(content: bytes, extension: str) -> str:
    """Route to the right extractor, correcting proven signature mismatches."""
    ext = extension.lower()
    if content.startswith(b"%PDF-") and ext != ".pdf":
        logger.warning(
            "[OneDrive] File extension %s disagrees with PDF signature; "
            "using the PDF extractor",
            ext or "[none]",
        )
        return _extract_text_from_pdf(content)
    if ext == ".pdf":
        return _extract_text_from_pdf(content)
    elif ext == ".docx":
        return _extract_text_from_docx(content)
    elif ext == ".doc":
        return _extract_text_from_legacy_doc(content)
    elif ext in (".txt", ".md", ".csv"):
        return content.decode("utf-8", errors="replace")
    elif ext == ".eml":
        return _extract_text_from_eml(content)
    elif ext == ".html":
        return _extract_text_from_html(content)
    elif ext == ".mpp":
        return _extract_text_from_mpp(content)
    elif ext == ".msg":
        return _extract_text_from_msg(content)
    elif ext == ".pptx":
        return _extract_text_from_pptx(content)
    elif ext == ".rtf":
        return _extract_text_from_rtf(content)
    elif ext == ".xls":
        return _extract_text_from_xls(content)
    elif ext in (".xlsx", ".xlsm"):
        return _extract_text_from_openxml_workbook(content)
    return ""


def _promote_to_project_documents(
    *,
    supabase_client,
    project_id: Optional[int],
    source_system: str,
    owner: str,
    folder_path: str,
    item: dict,
    item_id: str,
    name: str,
    storage_path: Optional[str],
    uploaded_by: str,
) -> None:
    if not project_id:
        return

    payload = project_document_payload_from_graph_item(
        project_id=project_id,
        source_system=source_system,
        owner=owner,
        folder_path=folder_path,
        item=item,
        item_id=item_id,
        name=name,
        storage_path=storage_path,
        uploaded_by=uploaded_by,
    )
    upsert_project_document_by_source(supabase_client, payload)


def sync_onedrive_folder(
    supabase_client,
    user_email: str,
    folder_path: str = "/",
    delta_token: Optional[str] = None,
) -> tuple[int, str]:
    """
    Sync files from a user's OneDrive folder. Returns (count_synced, new_delta_token).

    Args:
        supabase_client: Supabase service client
        user_email: The user whose OneDrive to sync
        folder_path: Folder path within OneDrive (default: root)
        delta_token: Previous delta token for incremental sync
    """
    graph = get_graph_client()
    if not graph.is_configured():
        logger.warning("[OneDrive] Microsoft Graph not configured — skipping")
        return 0, delta_token or ""

    user_id = user_email
    if folder_path == "/" or not folder_path:
        delta_path = f"/users/{user_id}/drive/root/delta"
    else:
        # Graph path format: /drive/root:/full/path/here:/delta
        # Do NOT replace inner slashes — only wrap the whole path with colons
        clean_path = folder_path.strip("/")
        delta_path = f"/users/{user_id}/drive/root:/{clean_path}:/delta"

    try:
        items, new_delta_token = graph.get_delta(delta_path, delta_token)
    except Exception as e:
        raise RuntimeError(f"OneDrive delta query failed for {user_email}{folder_path}: {e}") from e

    synced = 0
    processed_files = 0
    capped = False
    max_files = _bounded_int_env("GRAPH_INGEST_MAX_FILES_PER_FOLDER", 250, 1, 5000)
    for item in items:
        if "@removed" in item:
            continue

        # Skip folders
        if "folder" in item:
            continue

        name = item.get("name", "")
        size = int(item.get("size") or 0)

        # Check extension
        _, ext = os.path.splitext(name)
        ext = ext.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise RuntimeError(
                f"OneDrive file type cannot be silently excluded: {name} "
                f"({ext or 'no extension'}); add an extractor or an explicit governed "
                "exclusion before advancing the delta cursor"
            )

        max_file_size_bytes = _max_file_size_bytes()
        if size > max_file_size_bytes:
            raise RuntimeError(
                f"OneDrive file exceeds the {max_file_size_bytes}-byte ingestion limit: "
                f"{name} ({size} bytes); the delta cursor was not advanced"
            )

        item_id = item.get("id", "")
        doc_id = f"onedrive_{item_id}"
        storage_path = metadata_text_storage("onedrive", user_email, item_id, ext)

        # Check if already ingested
        existing = (
            supabase_client.from_("document_metadata")
            .select("id, project_id, business_area_id")
            .eq("id", doc_id)
            .limit(1)
            .execute()
        )
        rag_store = SupabaseRagStore(supabase_client)
        rag_replica = (
            rag_store.fetch_rag_document_metadata(doc_id) if existing.data else None
        )
        if existing.data and rag_replica:
            existing_doc = existing.data[0]
            if existing_doc.get("business_area_id"):
                assignment = AssignmentTarget(
                    project_id=None,
                    business_area_id=int(existing_doc["business_area_id"]),
                    method="existing_business_area",
                    confidence=1.0,
                )
            elif existing_doc.get("project_id"):
                assignment = _target_for_matched_project(
                    supabase_client,
                    project_id=int(existing_doc["project_id"]),
                    method="existing_project",
                )
            else:
                clean_content = _fetch_rag_document_text(doc_id)
                assignment = _assign_scope(
                    supabase_client,
                    item=item,
                    root_folder=folder_path,
                    title=name,
                    content=clean_content,
                    participants=[user_email],
                )
            _set_document_scope(supabase_client, doc_id, assignment)
            supabase_client.from_("document_metadata").update({
                "tags": ",".join(
                    ["onedrive", ext.lstrip("."), _assignment_tag(assignment)]
                ),
            }).eq("id", doc_id).execute()
            _promote_to_project_documents(
                supabase_client=supabase_client,
                project_id=assignment.project_id,
                source_system="onedrive",
                owner=user_email,
                folder_path=folder_path,
                item=item,
                item_id=item_id,
                name=name,
                storage_path=storage_path,
                uploaded_by=user_email,
            )
            continue
        if existing.data:
            logger.warning(
                "[OneDrive] App metadata exists without its RAG replica for %s; "
                "rehydrating the full document instead of applying a scope-only update",
                doc_id,
            )

        if processed_files >= max_files:
            capped = True
            logger.warning(
                "[OneDrive] File ingestion capped at %s new supported files for %s%s; "
                "preserving the prior delta cursor so remaining files replay next run",
                max_files,
                user_email,
                folder_path,
            )
            break
        processed_files += 1

        # Download file content
        download_url = item.get("@microsoft.graph.downloadUrl", "")
        if not download_url:
            try:
                file_data = graph.get(f"/users/{user_id}/drive/items/{item_id}")
                download_url = file_data.get("@microsoft.graph.downloadUrl", "")
            except Exception:
                continue

        if not download_url:
            continue

        try:
            raw_bytes = graph.download_bytes(download_url)
        except Exception as e:
            logger.warning(f"[OneDrive] Download failed for {name}: {e}")
            continue

        # Extract text (strip null bytes — postgres rejects \u0000 in text columns)
        text_content = _extract_text(raw_bytes, ext).replace("\x00", "")
        has_text = len(text_content.strip()) >= 50

        # Get file metadata
        modified = item.get("lastModifiedDateTime", datetime.now(timezone.utc).isoformat())
        web_url = item.get("webUrl", "")
        created_by = item.get("createdBy", {}).get("user", {}).get("displayName", user_email)

        # Only upload to storage when we have meaningful extracted text
        if has_text:
            try:
                storage_upload_with_retry(
                    supabase_client.storage.from_(DOCUMENT_BUCKET),
                    storage_path,
                    text_content.encode("utf-8"),
                    {"content-type": "text/plain", "upsert": "true"},
                )
            except Exception as e:
                logger.warning(f"[OneDrive] Storage upload failed for {name}: {e}")
        else:
            logger.debug(f"[OneDrive] Scanned/no-text file, saving metadata only: {name}")

        try:
            clean_content = text_content if has_text else ""
            # Always save metadata so the file appears in Files even when text extraction
            # failed (scanned PDFs, image-only documents, etc.)
            # Strip null bytes — PostgreSQL text columns reject \u0000
            assignment = _assign_scope(
                supabase_client,
                item=item,
                root_folder=folder_path,
                title=name,
                content=clean_content,
                participants=[created_by, user_email],
            )
            SupabaseRagStore(supabase_client).upsert_document_metadata({
                "id": doc_id,
                "title": name,
                "source": "microsoft_graph",
                "category": "document",
                "type": "document",
                "content": clean_content or None,
                "date": modified[:10] if modified else None,
                "url": web_url,
                "participants": ", ".join([created_by, user_email]),
                "status": "raw_ingested" if has_text else "no_text",
                "tags": ",".join(
                    ["onedrive", ext.lstrip("."), _assignment_tag(assignment)]
                ),
                "project_id": assignment.project_id,
                "business_area_id": assignment.business_area_id,
                "source_system": "onedrive",
                "source_item_id": item_id,
                "source_drive_id": graph_id_safe((item.get("parentReference") or {}).get("driveId")),
                "source_path": _item_source_path(item, folder_path, name),
                "source_web_url": web_url or None,
                "source_etag": item.get("eTag") or item.get("cTag"),
                "source_last_modified_at": modified,
                "source_size": size,
                "storage_bucket": DOCUMENT_BUCKET if has_text else None,
                "file_path": storage_path if has_text else None,
                "source_metadata": {
                    "graph_source": "onedrive",
                    "graph_owner": user_email,
                    "source_folder": folder_path,
                },
            })
            _promote_to_project_documents(
                supabase_client=supabase_client,
                project_id=assignment.project_id,
                source_system="onedrive",
                owner=user_email,
                folder_path=folder_path,
                item=item,
                item_id=item_id,
                name=name,
                storage_path=storage_path,
                uploaded_by=user_email,
            )
            synced += 1
            if (
                assignment.project_id is not None
                or assignment.business_area_id is not None
            ):
                logger.info(
                    "[OneDrive] Auto-assigned %s=%s for %s via %s (%.2f)",
                    "business_area_id"
                    if assignment.business_area_id is not None
                    else "project_id",
                    assignment.business_area_id or assignment.project_id,
                    item_id,
                    assignment.method,
                    assignment.confidence,
                )
        except Exception as e:
            logger.warning(f"[OneDrive] Failed to insert metadata for {name}: {e}")

    logger.info(f"[OneDrive] Synced {synced} files for {user_email}{folder_path} (processed={processed_files})")
    return synced, (delta_token or "") if capped else new_delta_token


def sync_sharepoint_folder(
    supabase_client,
    site_hostname: str,
    site_name: str,
    folder_path: str = "/",
    delta_token: Optional[str] = None,
    *,
    expected_project_number: str | None = None,
) -> SharePointSyncResult:
    """
    Sync files from a SharePoint site folder.

    A failed file keeps the previous delta cursor so the next scheduled pass
    automatically replays that delta batch. Existing documents are idempotent,
    so successfully processed files are not duplicated during recovery.

    Args:
        site_hostname: e.g. "alleato.sharepoint.com"
        site_name: e.g. "AlleatoGroup"
        folder_path: Folder path within the site drive (e.g. "/SOP" or "/")
        delta_token: Previous delta token for incremental sync
        expected_project_number: Exact YY-NNN number from governed discovery
    """
    graph = get_graph_client()
    if not graph.is_configured():
        logger.warning("[SharePoint] Microsoft Graph not configured — skipping")
        return SharePointSyncResult(0, delta_token or "")

    site_lookup = graph.get(f"/sites/{site_hostname}:/sites/{site_name}")
    site_id = str(site_lookup.get("id") or "").strip()
    if not site_id:
        logger.error("[SharePoint] Could not resolve site id for %s:/sites/%s", site_hostname, site_name)
        return SharePointSyncResult(0, delta_token or "")

    if folder_path == "/" or not folder_path:
        delta_path = f"/sites/{site_id}/drive/root/delta"
    else:
        clean_path = quote(folder_path.strip("/"), safe="/")
        delta_path = f"/sites/{site_id}/drive/root:/{clean_path}:/delta"

    try:
        get_delta_batch = getattr(graph, "get_delta_batch", None)
        if callable(get_delta_batch):
            delta_batch = get_delta_batch(
                delta_path,
                delta_token,
                max_pages=100,
                max_items=10000,
            )
            items = delta_batch.items
            new_delta_token = delta_batch.next_cursor
            inventory_complete = delta_batch.complete
        else:
            # Compatibility seam for test doubles and alternate Graph clients.
            items, new_delta_token = graph.get_delta(
                delta_path,
                delta_token,
                max_pages=100,
                max_items=10000,
            )
            inventory_complete = bool(new_delta_token)
    except Exception as e:
        raise RuntimeError(f"SharePoint delta query failed for {site_name}{folder_path}: {e}") from e
    if not inventory_complete and not new_delta_token:
        raise RuntimeError(
            "SharePoint delta inventory ended without either an @odata.deltaLink "
            "or an @odata.nextLink continuation cursor. The scope cursor was not "
            "advanced because completeness cannot be proven."
        )

    synced = 0
    failed = 0
    unchanged = 0
    excluded = 0
    excluded_by_extension: dict[str, int] = {}
    processed_files = 0
    capped = False
    max_files = _bounded_int_env("GRAPH_INGEST_MAX_FILES_PER_FOLDER", 250, 1, 5000)
    for item in items:
        if "@removed" in item or "folder" in item:
            continue

        name = item.get("name", "")
        size = item.get("size", 0)
        _, ext = os.path.splitext(name)
        ext = ext.lower()
        extension_label = ext or "[none]"
        if ext in GOVERNED_NON_TEXT_EXTENSIONS:
            excluded += 1
            excluded_by_extension[extension_label] = (
                excluded_by_extension.get(extension_label, 0) + 1
            )
            continue
        if ext not in SUPPORTED_EXTENSIONS:
            failed += 1
            logger.error(
                "[SharePoint] %s file type cannot be silently excluded: %s (%s). "
                "Add a complete extractor or an explicit governed non-text classification "
                "before advancing the delta cursor.",
                (
                    "Text-bearing unsupported"
                    if ext in TEXT_BEARING_UNSUPPORTED_EXTENSIONS
                    else "Unknown"
                ),
                name,
                extension_label,
            )
            continue
        max_file_size_bytes = _max_file_size_bytes()
        route_pdf_to_streamed_ocr = ext == ".pdf"
        if route_pdf_to_streamed_ocr and size > _STREAMED_PDF_MAX_BYTES:
            failed += 1
            logger.error(
                "[SharePoint] PDF exceeds the governed memory-safe streamed "
                "OCR limit of %d bytes: %s (%d bytes)",
                _STREAMED_PDF_MAX_BYTES,
                name,
                size,
            )
            continue
        if not route_pdf_to_streamed_ocr and size > max_file_size_bytes:
            failed += 1
            logger.error(
                "[SharePoint] Supported file exceeds the %d-byte ingestion limit: %s (%d bytes)",
                max_file_size_bytes,
                name,
                size,
            )
            continue

        item_id = item.get("id", "")
        doc_id = f"sharepoint_{item_id}"
        storage_path = metadata_text_storage("sharepoint", site_name, item_id, ext)
        source_etag = item.get("eTag") or item.get("cTag")
        modified = item.get(
            "lastModifiedDateTime",
            datetime.now(timezone.utc).isoformat(),
        )

        existing = (
            supabase_client.from_("document_metadata")
            .select(
                "id,project_id,business_area_id,source_etag,source_last_modified_at,status"
            )
            .eq("id", doc_id)
            .limit(1)
            .execute()
        )
        existing_doc = existing.data[0] if existing.data else None
        rag_store = SupabaseRagStore(supabase_client)
        rag_replica = (
            rag_store.fetch_rag_document_metadata(doc_id) if existing_doc else None
        )
        if existing_doc:
            same_revision = bool(rag_replica) and (
                bool(
                    source_etag
                    and str(existing_doc.get("source_etag") or "") == str(source_etag)
                )
                or bool(
                    not source_etag
                    and modified
                    and str(existing_doc.get("source_last_modified_at") or "")
                    == str(modified)
                )
            )
            if not rag_replica:
                logger.warning(
                    "[SharePoint] App metadata exists without its RAG replica for %s; "
                    "rehydrating the full document instead of applying a scope-only update",
                    doc_id,
                )
            elif not same_revision:
                logger.info(
                    "[SharePoint] Source revision changed for %s; refreshing full text",
                    name,
                )
                rag_store.invalidate_document_content(
                    doc_id,
                    parsing_status="source_revision_changed",
                    embedding_status="pending",
                    reason=(
                        "SharePoint eTag or last-modified revision changed; "
                        "previous chunks cannot remain searchable"
                    ),
                )
            else:
                unchanged += 1
                if existing_doc.get("business_area_id"):
                    assignment = AssignmentTarget(
                        project_id=None,
                        business_area_id=int(existing_doc["business_area_id"]),
                        method="existing_business_area",
                        confidence=1.0,
                    )
                elif existing_doc.get("project_id"):
                    assignment = _target_for_matched_project(
                        supabase_client,
                        project_id=int(existing_doc["project_id"]),
                        method="existing_project",
                    )
                else:
                    clean_content = _fetch_rag_document_text(doc_id)
                    assignment = _assign_scope(
                        supabase_client,
                        item=item,
                        root_folder=folder_path,
                        title=name,
                        content=clean_content,
                        participants=[site_name],
                        expected_project_number=expected_project_number,
                    )
                _set_document_scope(supabase_client, doc_id, assignment)
                supabase_client.from_("document_metadata").update(
                    {
                        "tags": ",".join(
                            [
                                "sharepoint",
                                site_name.lower(),
                                ext.lstrip("."),
                                _assignment_tag(assignment),
                            ]
                        )
                    }
                ).eq("id", doc_id).execute()
                _promote_to_project_documents(
                    supabase_client=supabase_client,
                    project_id=assignment.project_id,
                    source_system="sharepoint",
                    owner=site_name,
                    folder_path=folder_path,
                    item=item,
                    item_id=item_id,
                    name=name,
                    storage_path=storage_path,
                    uploaded_by="SharePoint sync",
                )
                continue

        if processed_files >= max_files:
            capped = True
            logger.warning(
                "[SharePoint] File ingestion capped at %s new or changed supported "
                "files for %s%s; preserving the prior delta cursor so remaining "
                "files replay next run",
                max_files,
                site_name,
                folder_path,
            )
            break
        processed_files += 1

        if existing_doc and existing_doc.get("business_area_id"):
            assignment = AssignmentTarget(
                project_id=None,
                business_area_id=int(existing_doc["business_area_id"]),
                method="existing_business_area",
                confidence=1.0,
            )
        elif existing_doc and existing_doc.get("project_id"):
            assignment = _target_for_matched_project(
                supabase_client,
                project_id=int(existing_doc["project_id"]),
                method="existing_project",
            )
        else:
            assignment = AssignmentTarget(
                project_id=None,
                business_area_id=None,
                method="unassigned",
                confidence=0.0,
            )

        web_url = item.get("webUrl", "")
        raw_bytes: bytes | None = None
        if route_pdf_to_streamed_ocr:
            if not web_url:
                failed += 1
                logger.error(
                    "[SharePoint] PDF cannot be queued for memory-safe Azure "
                    "extraction because Graph returned no webUrl: %s",
                    name,
                )
                continue
            text_content = ""
        else:
            download_url = item.get("@microsoft.graph.downloadUrl", "")
            if not download_url:
                try:
                    file_data = graph.get(f"/sites/{site_id}/drive/items/{item_id}")
                    download_url = file_data.get("@microsoft.graph.downloadUrl", "")
                except Exception as exc:
                    failed += 1
                    logger.error("[SharePoint] Could not get a download URL for %s: %s", name, exc)
                    continue

            if not download_url:
                failed += 1
                logger.error("[SharePoint] No download URL returned for %s", name)
                continue

            try:
                raw_bytes = graph.download_bytes(download_url)
            except Exception as e:
                failed += 1
                logger.error("[SharePoint] Download failed for %s: %s", name, e)
                continue
            try:
                text_content = _extract_text(raw_bytes, ext).replace("\x00", "")
            except Exception as exc:
                failed += 1
                logger.error(
                    "[SharePoint] Complete text extraction failed for %s: %s",
                    name,
                    exc,
                )
                continue

        has_text = len(text_content.strip()) >= 50
        embedded_image_names = (
            _docx_embedded_image_names(raw_bytes)
            if ext == ".docx" and not has_text and raw_bytes is not None
            else []
        )
        requires_ocr = route_pdf_to_streamed_ocr or bool(embedded_image_names)
        if not has_text and not requires_ocr:
            failed += 1
            logger.error(
                "[SharePoint] Text extraction produced insufficient content for supported file %s (%s)",
                name,
                ext,
            )
            continue
        if requires_ocr:
            rag_store = rag_store or SupabaseRagStore(supabase_client)
            rag_store.invalidate_document_content(
                doc_id,
                parsing_status="no_text",
                embedding_status="pending_ocr",
                reason=(
                    "The current SharePoint PDF revision contains no extractable "
                    "text and is awaiting complete OCR"
                ),
            )

        created_by = item.get("createdBy", {}).get("user", {}).get("displayName", site_name)

        if has_text:
            try:
                storage_upload_with_retry(
                    supabase_client.storage.from_(DOCUMENT_BUCKET),
                    storage_path,
                    text_content.encode("utf-8"),
                    {"content-type": "text/plain", "upsert": "true"},
                )
            except Exception as e:
                failed += 1
                logger.error("[SharePoint] Storage upload failed for %s: %s", name, e)
                continue

        try:
            clean_content = text_content
            if (
                assignment.project_id is None
                and assignment.business_area_id is None
            ):
                assignment = _assign_scope(
                    supabase_client,
                    item=item,
                    root_folder=folder_path,
                    title=name,
                    content=clean_content,
                    participants=[created_by, site_name],
                    expected_project_number=expected_project_number,
                )
            content_hash = (
                hashlib.sha256(clean_content.encode("utf-8")).hexdigest()
                if has_text
                else None
            )
            rag_store = rag_store or SupabaseRagStore(supabase_client)
            rag_store.upsert_document_metadata({
                "id": doc_id,
                "title": name,
                "source": "microsoft_graph",
                "category": "document",
                "type": "document",
                "content": clean_content if has_text else None,
                "date": modified[:10] if modified else None,
                "url": web_url,
                "participants": ", ".join([created_by, site_name]),
                "status": "raw_ingested" if has_text else "no_text",
                "parsing_status": "raw_ingested" if has_text else "no_text",
                "embedding_status": "pending" if has_text else None,
                "tags": ",".join(
                    [
                        "sharepoint",
                        site_name.lower(),
                        ext.lstrip("."),
                        _assignment_tag(assignment),
                    ]
                ),
                "project_id": assignment.project_id,
                "business_area_id": assignment.business_area_id,
                "source_system": "sharepoint",
                "source_item_id": item_id,
                "source_drive_id": graph_id_safe((item.get("parentReference") or {}).get("driveId")),
                "source_site_id": graph_id_safe((item.get("parentReference") or {}).get("siteId")),
                "source_path": _item_source_path(item, folder_path, name),
                "source_web_url": web_url or None,
                "source_etag": source_etag,
                "source_last_modified_at": modified,
                "source_size": size,
                "content_hash": content_hash,
                "storage_bucket": DOCUMENT_BUCKET if has_text else None,
                "file_path": storage_path if has_text else None,
                "last_synced_at": datetime.now(timezone.utc).isoformat(),
                "source_metadata": {
                    "graph_source": "sharepoint",
                    "graph_owner": site_name,
                    "source_folder": folder_path,
                    "project_number": expected_project_number,
                    "content_hash": content_hash,
                    "ocr_required": requires_ocr,
                    "ocr_transport": (
                        "streamed_pdf"
                        if route_pdf_to_streamed_ocr
                        else "openxml_embedded_images"
                        if embedded_image_names
                        else None
                    ),
                    "ocr_embedded_images": embedded_image_names,
                },
            })
            _promote_to_project_documents(
                supabase_client=supabase_client,
                project_id=assignment.project_id,
                source_system="sharepoint",
                owner=site_name,
                folder_path=folder_path,
                item=item,
                item_id=item_id,
                name=name,
                storage_path=storage_path if has_text else None,
                uploaded_by="SharePoint sync",
            )
            synced += 1
        except Exception as e:
            failed += 1
            logger.error("[SharePoint] Failed to insert metadata for %s: %s", name, e)

    retry_required = failed > 0 or capped or not inventory_complete
    if failed > 0:
        retry_reason = (
            f"{failed} SharePoint file(s) failed; the prior delta cursor was preserved for automatic retry."
        )
    elif capped:
        retry_reason = (
            f"SharePoint ingestion reached its {max_files}-new-file safety limit; "
            "the prior delta cursor was preserved so remaining files replay next run."
        )
    elif not inventory_complete:
        retry_reason = (
            "SharePoint delta inventory reached its 100-page/10,000-item safety "
            "limit; the continuation cursor was saved so the next run resumes "
            "at the next unread Graph page."
        )
    else:
        retry_reason = None
    if failed > 0 or capped:
        saved_delta_token = delta_token or ""
    else:
        saved_delta_token = new_delta_token
    if retry_required:
        logger.error(
            "[SharePoint] Recovery required in %s%s: %s",
            site_name,
            folder_path,
            retry_reason,
        )
    logger.info(
        "[SharePoint] Synced %d files from %s%s "
        "(processed=%d, unchanged=%d, excluded=%d, failed=%d)",
        synced,
        site_name,
        folder_path,
        processed_files,
        unchanged,
        excluded,
        failed,
    )
    return SharePointSyncResult(
        items_synced=synced,
        delta_token=saved_delta_token,
        items_failed=failed,
        items_unchanged=unchanged,
        items_excluded=excluded,
        excluded_by_extension=excluded_by_extension,
        retry_required=retry_required,
        retry_reason=retry_reason,
        inventory_complete=inventory_complete,
    )
