import importlib
import io
import sys
import types
from datetime import date
from types import SimpleNamespace
from zipfile import ZipFile

import pytest
from src.services.ingestion.project_assignment import AssignmentTarget
from src.services.integrations.microsoft_graph import onedrive


class _Result:
    def __init__(self, data=None):
        self.data = data or []


class _StorageBucket:
    def __init__(self):
        self.uploads = []

    def upload(self, path, content, options):
        self.uploads.append((path, content, options))
        return _Result([{"path": path}])


class _Storage:
    def __init__(self):
        self.bucket = _StorageBucket()

    def from_(self, _name):
        return self.bucket


class _Table:
    def __init__(self, store, name):
        self.store = store
        self.name = name
        self.filters = {}
        self.payload = None
        self.action = "select"

    def select(self, *_args, **_kwargs):
        return self

    def eq(self, key, value):
        self.filters[key] = value
        return self

    def is_(self, key, value):
        self.filters[key] = value
        return self

    def limit(self, *_args, **_kwargs):
        return self

    def insert(self, payload):
        self.action = "insert"
        self.payload = payload
        return self

    def update(self, payload):
        self.action = "update"
        self.payload = payload
        return self

    def upsert(self, payload):
        self.action = "upsert"
        self.payload = payload
        return self

    def execute(self):
        rows = self.store.setdefault(self.name, [])
        matches = [
            row
            for row in rows
            if all(row.get(key) == value for key, value in self.filters.items())
        ]

        if self.action == "insert":
            row = {"id": len(rows) + 1, **self.payload}
            rows.append(row)
            return _Result([row])

        if self.action == "update":
            for row in matches:
                row.update(self.payload)
            return _Result(matches)

        if self.action == "upsert":
            payload = self.payload
            existing = next(
                (row for row in rows if row.get("id") == payload.get("id")), None
            )
            if existing:
                existing.update(payload)
                return _Result([existing])
            rows.append(dict(payload))
            return _Result([payload])

        return _Result(matches)


class _Supabase:
    def __init__(self):
        self.store = {}
        self.storage = _Storage()

    def from_(self, name):
        return _Table(self.store, name)

    def table(self, name):
        return self.from_(name)


class _OfflineRagStore:
    def __init__(self, supabase):
        self.supabase = supabase

    def upsert_document_metadata(self, payload):
        catalog = {
            key: value
            for key, value in payload.items()
            if key not in {"content", "raw_text", "summary_embedding"}
        }
        self.supabase.from_("document_metadata").upsert(catalog).execute()
        self.supabase.from_("rag_document_metadata").upsert(dict(payload)).execute()
        return payload

    def fetch_rag_document_metadata(self, document_id):
        rows = self.supabase.store.get("rag_document_metadata", [])
        return next((row for row in rows if row.get("id") == document_id), None)


@pytest.fixture(autouse=True)
def _use_offline_rag_store(monkeypatch):
    monkeypatch.setattr(onedrive, "SupabaseRagStore", _OfflineRagStore)


class _Graph:
    def __init__(self, items):
        self.items = items
        self.downloads = []

    def is_configured(self):
        return True

    def get_delta(self, path, delta_token):
        assert path == "/users/pm@example.com/drive/root:/Projects:/delta"
        assert delta_token is None
        return self.items, "next-token"

    def download_bytes(self, download_url):
        self.downloads.append(download_url)
        return b"This is enough project scope text to pass extraction and assign a project."


class _FailingDeltaGraph(_Graph):
    def __init__(self):
        super().__init__([])

    def get_delta(self, path, delta_token):
        raise RuntimeError("404 Not Found")


def _workbook_bytes(*, large=False, cached_formula=False):
    # The shared backend conftest stubs NumPy before application import. Reload
    # the real dependency before importing openpyxl so this focused extractor
    # test exercises the same dependency path as production.
    if not isinstance(sys.modules.get("numpy"), types.ModuleType):
        sys.modules.pop("numpy", None)
        for module_name in list(sys.modules):
            if module_name == "openpyxl" or module_name.startswith("openpyxl."):
                sys.modules.pop(module_name, None)
        importlib.import_module("numpy")
    from openpyxl import Workbook

    workbook = Workbook()
    summary = workbook.active
    summary.title = "Summary"
    summary["A1"] = "Union Collective estimate"
    summary["A2"] = 125000
    summary["A3"] = 75000
    summary["A4"] = "=SUM(A2:A3)"

    details = workbook.create_sheet("Scope Detail")
    details["B2"] = "Fire sprinkler material takeoff"
    details["C2"] = date(2026, 7, 22)
    if large:
        for row_index in range(1, 81):
            details.cell(
                row=row_index + 2, column=1, value=f"Line {row_index}: " + ("x" * 900)
            )

    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    workbook_content = output.getvalue()
    if not cached_formula:
        return workbook_content

    # openpyxl does not calculate formulas. Inject the cached value that Excel
    # persists beside the formula so the dual-load extractor can prove it keeps
    # both the expression and the last calculated result.
    patched = io.BytesIO()
    with (
        ZipFile(io.BytesIO(workbook_content), "r") as source,
        ZipFile(patched, "w") as target,
    ):
        for entry in source.infolist():
            content = source.read(entry.filename)
            if entry.filename == "xl/worksheets/sheet1.xml":
                content = content.replace(
                    b"<f>SUM(A2:A3)</f><v></v>",
                    b"<f>SUM(A2:A3)</f><v>200000</v>",
                )
            target.writestr(entry, content)
    return patched.getvalue()


def _drive_item():
    return {
        "id": "drive-item-1",
        "name": "Scope.txt",
        "size": 72,
        "@microsoft.graph.downloadUrl": "https://download.example/scope",
        "webUrl": "https://microsoft.example/scope",
        "lastModifiedDateTime": "2026-05-12T10:00:00Z",
        "eTag": "etag-1",
        "parentReference": {"driveId": "drive-1"},
        "createdBy": {"user": {"displayName": "Project Manager"}},
    }


def _drive_item_with_id(item_id):
    item = _drive_item()
    item["id"] = item_id
    item["name"] = f"Scope {item_id}.txt"
    item["@microsoft.graph.downloadUrl"] = f"https://download.example/{item_id}"
    return item


def test_onedrive_sync_promotes_assigned_file_to_project_documents(monkeypatch):
    supabase = _Supabase()
    graph = _Graph([_drive_item()])
    monkeypatch.setattr(onedrive, "get_graph_client", lambda: graph)
    monkeypatch.setattr(
        onedrive,
        "infer_assignment_target",
        lambda *_args, **_kwargs: AssignmentTarget(
            project_id=25125,
            business_area_id=None,
            method="project_number",
            confidence=0.95,
        ),
    )

    count, token = onedrive.sync_onedrive_folder(
        supabase, "pm@example.com", "/Projects"
    )

    assert count == 1
    assert token == "next-token"
    assert graph.downloads == ["https://download.example/scope"]
    assert (
        supabase.storage.bucket.uploads[0][0]
        == "onedrive/pm@example.com/drive-item-1.txt.txt"
    )

    metadata = supabase.store["document_metadata"][0]
    assert metadata["id"] == "onedrive_drive-item-1"
    assert metadata["project_id"] == 25125
    assert metadata["source_system"] == "onedrive"
    assert metadata["file_path"] == "onedrive/pm@example.com/drive-item-1.txt.txt"

    project_doc = supabase.store["project_documents"][0]
    assert project_doc["project_id"] == 25125
    assert project_doc["source_system"] == "onedrive"
    assert project_doc["source_item_id"] == "drive-item-1"
    assert project_doc["file_url"] == "https://microsoft.example/scope"
    assert project_doc["storage_path"] is None
    assert project_doc["source_metadata"]["text_storage_path"] == metadata["file_path"]


def test_existing_onedrive_metadata_still_promotes_to_project_documents(monkeypatch):
    supabase = _Supabase()
    supabase.store["document_metadata"] = [
        {
            "id": "onedrive_drive-item-1",
            "project_id": 25125,
            "content": "Existing extracted text",
        }
    ]
    supabase.store["rag_document_metadata"] = [
        {
            "id": "onedrive_drive-item-1",
            "content": "Existing extracted text",
        }
    ]
    graph = _Graph([_drive_item()])
    monkeypatch.setattr(onedrive, "get_graph_client", lambda: graph)
    monkeypatch.setattr(
        onedrive,
        "_target_for_matched_project",
        lambda *_args, **_kwargs: AssignmentTarget(
            project_id=25125,
            business_area_id=None,
            method="existing_project",
            confidence=1.0,
        ),
    )

    count, _token = onedrive.sync_onedrive_folder(
        supabase, "pm@example.com", "/Projects"
    )

    assert count == 0
    assert graph.downloads == []
    project_doc = supabase.store["project_documents"][0]
    assert project_doc["project_id"] == 25125
    assert project_doc["source_system"] == "onedrive"
    assert project_doc["source_item_id"] == "drive-item-1"


def test_existing_onedrive_metadata_without_rag_replica_rehydrates(monkeypatch):
    supabase = _Supabase()
    supabase.store["document_metadata"] = [
        {
            "id": "onedrive_drive-item-1",
            "project_id": 25125,
        }
    ]
    graph = _Graph([_drive_item()])
    monkeypatch.setattr(onedrive, "get_graph_client", lambda: graph)
    monkeypatch.setattr(
        onedrive,
        "_target_for_matched_project",
        lambda *_args, **_kwargs: AssignmentTarget(
            project_id=25125,
            business_area_id=None,
            method="existing_project",
            confidence=1.0,
        ),
    )

    count, token = onedrive.sync_onedrive_folder(
        supabase,
        "pm@example.com",
        "/Projects",
    )

    assert count == 1
    assert token == "next-token"
    assert graph.downloads == ["https://download.example/scope"]
    assert supabase.store["rag_document_metadata"][0]["id"] == ("onedrive_drive-item-1")
    assert (
        "enough project scope text"
        in (supabase.store["rag_document_metadata"][0]["content"])
    )


def test_onedrive_sync_caps_supported_files_per_folder(monkeypatch, caplog):
    supabase = _Supabase()
    graph = _Graph([_drive_item_with_id("a"), _drive_item_with_id("b")])
    monkeypatch.setenv("GRAPH_INGEST_MAX_FILES_PER_FOLDER", "1")
    monkeypatch.setattr(onedrive, "get_graph_client", lambda: graph)
    monkeypatch.setattr(
        onedrive,
        "infer_assignment_target",
        lambda *_args, **_kwargs: AssignmentTarget(
            project_id=25125,
            business_area_id=None,
            method="project_number",
            confidence=0.95,
        ),
    )

    count, token = onedrive.sync_onedrive_folder(
        supabase, "pm@example.com", "/Projects"
    )

    assert count == 1
    assert token == ""
    assert graph.downloads == ["https://download.example/a"]
    assert "File ingestion capped at 1 new supported files" in caplog.text

    replay_count, replay_token = onedrive.sync_onedrive_folder(
        supabase, "pm@example.com", "/Projects"
    )

    assert replay_count == 1
    assert replay_token == "next-token"
    assert graph.downloads == [
        "https://download.example/a",
        "https://download.example/b",
    ]


def test_pdf_optional_dependency_warning_logs_once(monkeypatch, caplog):
    real_import = __import__

    def fake_import(name, *args, **kwargs):
        if name in {"pypdf", "fitz"}:
            raise ImportError(name)
        return real_import(name, *args, **kwargs)

    onedrive._WARNED_OPTIONAL_DEPENDENCIES.clear()
    monkeypatch.setattr("builtins.__import__", fake_import)

    assert onedrive._extract_text_from_pdf(b"%PDF no text") == ""
    assert onedrive._extract_text_from_pdf(b"%PDF no text") == ""

    assert caplog.text.count("pypdf not installed") == 1
    assert caplog.text.count("PyMuPDF not installed") == 1


def test_pdf_extraction_reads_pages_after_the_previous_fifty_page_cap(monkeypatch):
    pages = [
        SimpleNamespace(
            extract_text=lambda index=index: f"Page {index} complete source text"
        )
        for index in range(1, 52)
    ]
    fake_pypdf = SimpleNamespace(PdfReader=lambda _stream: SimpleNamespace(pages=pages))
    monkeypatch.setitem(sys.modules, "pypdf", fake_pypdf)

    text = onedrive._extract_text_from_pdf(b"fake-pdf")

    assert "Page 1 complete source text" in text
    assert "Page 51 complete source text" in text


def test_docx_extraction_includes_tables_headers_and_footers():
    import docx

    document = docx.Document()
    document.add_paragraph("Executive proposal narrative")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Base bid"
    table.cell(0, 1).text = "$250,000"
    document.sections[0].header.paragraphs[0].text = "Union Collective"
    document.sections[0].footer.paragraphs[0].text = "Confidential estimate"
    output = io.BytesIO()
    document.save(output)

    text = onedrive._extract_text_from_docx(output.getvalue())

    assert "Executive proposal narrative" in text
    assert "R1C1: Base bid" in text
    assert "R1C2: $250,000" in text
    assert "[Header]\nUnion Collective" in text
    assert "[Footer]\nConfidential estimate" in text


@pytest.mark.parametrize("extension", [".xlsx", ".xlsm"])
def test_openxml_workbooks_materialize_all_sheets_cells_and_formulas(extension):
    text = onedrive._extract_text(_workbook_bytes(cached_formula=True), extension)

    assert "[Worksheet: Summary]" in text
    assert "A1: Union Collective estimate" in text
    assert "A2: 125000" in text
    assert "A4 formula: =SUM(A2:A3) | cached value: 200000" in text
    assert "[Worksheet: Scope Detail]" in text
    assert "B2: Fire sprinkler material takeoff" in text
    assert "C2: 2026-07-22T00:00:00" in text


def test_openxml_workbook_extraction_does_not_truncate_large_materialized_text():
    text = onedrive._extract_text(_workbook_bytes(large=True), ".xlsx")

    assert len(text) > 50000
    assert "Line 1:" in text
    assert "Line 80:" in text


def test_openxml_workbook_falls_back_when_named_styles_are_invalid():
    workbook = _workbook_bytes(cached_formula=True)
    source = io.BytesIO(workbook)
    output = io.BytesIO()
    with ZipFile(source) as input_archive, ZipFile(output, "w") as output_archive:
        for name in input_archive.namelist():
            payload = input_archive.read(name)
            if name == "xl/styles.xml":
                payload = payload.replace(
                    b'<cellStyle name="Normal" xfId="0" builtinId="0"/>',
                    b'<cellStyle name="Normal" xfId="999" builtinId="0"/>',
                )
            output_archive.writestr(name, payload)

    text = onedrive._extract_text(output.getvalue(), ".xlsx")

    assert "[Worksheet: Summary]" in text
    assert "A1: Union Collective estimate" in text
    assert "A4 formula: =SUM(A2:A3) | cached value: 200000" in text
    assert "[Worksheet: Scope Detail]" in text
    assert "B2: Fire sprinkler material takeoff" in text


def test_mislabeled_docx_with_pdf_signature_uses_pdf_extractor(monkeypatch):
    calls = []
    monkeypatch.setattr(
        onedrive,
        "_extract_text_from_pdf",
        lambda content: calls.append(content) or "Complete proposal text",
    )

    content = b"%PDF-1.4\nmislabeled proposal"
    extracted = onedrive._extract_text(content, ".docx")

    assert extracted == "Complete proposal text"
    assert calls == [content]


def test_workbook_extensions_are_part_of_the_canonical_graph_contract():
    assert {".xlsx", ".xlsm"}.issubset(onedrive.SUPPORTED_EXTENSIONS)


def test_onedrive_delta_failure_raises_for_orchestrator(monkeypatch):
    monkeypatch.setattr(onedrive, "get_graph_client", lambda: _FailingDeltaGraph())

    with pytest.raises(RuntimeError, match="OneDrive delta query failed"):
        onedrive.sync_onedrive_folder(_Supabase(), "pm@example.com", "/Projects")
